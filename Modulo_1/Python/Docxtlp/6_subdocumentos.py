from docxtpl import DocxTemplate
from docx import Document

# 1. Crear plantilla master
doc = Document()
doc.add_paragraph('Inicio Master.')
doc.add_paragraph('{{ subdocumento }}')
doc.add_paragraph('Fin Master.')
doc.save('plantilla_master.docx')

# 2. Crear subdocumento
sub = Document()
sub.add_heading('Sub-documento', level=2)
sub.add_paragraph('Contenido importado.')
sub.save('anexo.docx')

# 3. Renderizar
doc_tpl = DocxTemplate("plantilla_master.docx")
sub_doc = doc_tpl.new_subdoc("anexo.docx")
context = {'subdocumento': sub_doc}

doc_tpl.render(context)
doc_tpl.save("resultado_subdoc.docx")
print("Generado: resultado_subdoc.docx")
