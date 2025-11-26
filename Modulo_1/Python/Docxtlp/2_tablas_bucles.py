from docxtpl import DocxTemplate
from docx import Document

# 1. Crear plantilla con tabla
doc = Document()
doc.add_heading('Reporte de Notas', 0)
table = doc.add_table(rows=2, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Estudiante'
hdr_cells[1].text = 'Materia'
hdr_cells[2].text = 'Nota'

# Fila dinámica
row_cells = table.rows[1].cells
row_cells[0].text = '{% for a in alumnos %}{{ a.nombre }}'
row_cells[1].text = '{{ a.materia }}'
row_cells[2].text = '{{ a.nota }}{% endfor %}'
doc.save('plantilla_tabla.docx')

# 2. Renderizar
doc_tpl = DocxTemplate("plantilla_tabla.docx")

lista_alumnos = [
    {'nombre': 'Juan', 'materia': 'Matemáticas', 'nota': 9.5},
    {'nombre': 'Maria', 'materia': 'Física', 'nota': 8.0},
    {'nombre': 'Pedro', 'materia': 'Química', 'nota': 6.5},
]

context = {'alumnos': lista_alumnos}
doc_tpl.render(context)
doc_tpl.save("resultado_tabla.docx")
print("Generado: resultado_tabla.docx")
