from docxtpl import DocxTemplate
from docx import Document

# 1. Crear plantilla
doc = Document()
doc.add_paragraph('Resultado del examen:')
doc.add_paragraph('{% if nota >= 60 %}APROBADO{% else %}REPROBADO{% endif %}')
doc.save('plantilla_condicional.docx')

# 2. Renderizar Aprobado
doc_tpl = DocxTemplate("plantilla_condicional.docx")
doc_tpl.render({'nota': 85})
doc_tpl.save("resultado_aprobado.docx")

# 3. Renderizar Reprobado
doc_tpl = DocxTemplate("plantilla_condicional.docx")
doc_tpl.render({'nota': 45})
doc_tpl.save("resultado_reprobado.docx")

print("Generados: resultado_aprobado.docx y resultado_reprobado.docx")
