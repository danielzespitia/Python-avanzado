from docxtpl import DocxTemplate, InlineImage
from docx import Document
from docx.shared import Mm
from PIL import Image
import os

# 1. Crear imagen dummy si no existe
if not os.path.exists('test_image.png'):
    img = Image.new('RGB', (100, 100), color = 'blue')
    img.save('test_image.png')

# 2. Crear plantilla
doc = Document()
doc.add_paragraph('Imagen dinámica:')
doc.add_paragraph('{{ imagen_dinamica }}')
doc.save('plantilla_imagen.docx')

# 3. Renderizar
doc_tpl = DocxTemplate("plantilla_imagen.docx")
imagen = InlineImage(doc_tpl, 'test_image.png', width=Mm(50))
context = {'imagen_dinamica': imagen}

doc_tpl.render(context)
doc_tpl.save("resultado_imagen.docx")
print("Generado: resultado_imagen.docx")
