from docxtpl import DocxTemplate
from docx import Document

# 1. Crear plantilla temporal (normalmente esto se hace en Word)
doc = Document()
doc.add_heading('Plantilla de Prueba', 0)
p = doc.add_paragraph('Hola ')
p.add_run('{{ nombre }}').bold = True
p.add_run(',')
doc.add_paragraph('Bienvenido al curso de {{ curso }}.')
doc.add_paragraph('Fecha: {{ fecha }}')
doc.save('plantilla_basica.docx')

# 2. Cargar plantilla
doc_tpl = DocxTemplate("plantilla_basica.docx")

# 3. Crear contexto
context = {
    'nombre': 'Daniel Espitia',
    'curso': 'Python Avanzado',
    'fecha': '25 de Noviembre de 2025'
}

# 4. Renderizar y guardar
doc_tpl.render(context)
doc_tpl.save("resultado_basico.docx")
print("Generado: resultado_basico.docx")
