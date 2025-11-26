from docx2pdf import convert
import os
from docx import Document

# 1. Crear archivo dummy si no existe
if not os.path.exists('documento_prueba.docx'):
    doc = Document()
    doc.add_heading('Prueba de Conversión', 0)
    doc.add_paragraph('Hola, esto es un PDF generado desde Python.')
    doc.save('documento_prueba.docx')

# 2. Convertir
print("Convirtiendo documento_prueba.docx a PDF...")
convert("documento_prueba.docx")
print("¡Hecho! Revisa documento_prueba.pdf")
